#!/usr/bin/env python3
"""Read a tiny deterministic sample from highest-ranked local source families.

Inputs and outputs are local operational state only. The output may contain
source locators and extracted text fragments, so it must never be versioned.
No source file is modified and no external service is contacted.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import re
from datetime import UTC, datetime
from pathlib import Path


ALLOWED = {".pdf", ".docx", ".xlsx"}
KEYWORDS = [
    "progress", "planned", "actual", "schedule", "delay", "risk", "issue", "milestone",
    "procurement", "engineering", "construction", "action", "management", "reporting period",
    "پیشرفت", "برنامه", "واقعی", "زمانبندی", "تاخیر", "تأخیر", "ریسک", "مسئله", "اقدام",
    "تدارکات", "مهندسی", "ساخت", "مدیریت", "گزارش",
]


def runtime(name: str) -> Path:
    return Path(os.environ["LOCALAPPDATA"]) / "EnterpriseAI" / "runtime" / name


def extract_pdf(path: Path) -> dict:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "")[:10000] for page in reader.pages[:8])
    return {"kind": "pdf", "pages_sampled": min(8, len(reader.pages)), "text": text}


def extract_docx(path: Path) -> dict:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs[:80] if paragraph.text]
        tables = [" | ".join(cell.text for row in table.rows[:8] for cell in row.cells[:12]) for table in doc.tables[:8]]
        return {"kind": "docx", "paragraphs_sampled": len(paragraphs), "tables_sampled": len(tables), "text": "\n".join(paragraphs + tables)[:60000]}
    except ModuleNotFoundError:
        import zipfile
        from xml.etree import ElementTree
        with zipfile.ZipFile(path) as archive:
            root = ElementTree.fromstring(archive.read("word/document.xml"))
        text = "\n".join("".join(node.itertext()) for node in root.findall(".//{*}p"))
        return {"kind": "docx", "paragraphs_sampled": len(root.findall(".//{*}p")), "tables_sampled": len(root.findall(".//{*}tbl")), "text": text[:60000]}


def extract_xlsx(path: Path) -> dict:
    from openpyxl import load_workbook
    book = load_workbook(path, read_only=True, data_only=False)
    chunks = []
    sheet_summaries = []
    for sheet in book.worksheets[:8]:
        values = []
        for row in sheet.iter_rows(max_row=80, max_col=24, values_only=True):
            rendered = " | ".join(str(value) for value in row if value not in (None, ""))
            if rendered:
                values.append(rendered)
        sheet_summaries.append({"sheet": sheet.title, "sampled_nonempty_rows": len(values)})
        chunks.extend(values)
    return {"kind": "xlsx", "sheets_sampled": sheet_summaries, "text": "\n".join(chunks)[:60000]}


def extract(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".xlsx":
        return extract_xlsx(path)
    raise ValueError("unsupported format")


def evidence_signals(text: str) -> list[str]:
    lowered = text.casefold()
    return [term for term in KEYWORDS if term.casefold() in lowered]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--ranking", type=Path, default=runtime("st1-040-source-family-ranking.json"))
    parser.add_argument("--family-manifest", type=Path, help="runtime-local single family manifest from discovery")
    parser.add_argument("--output", type=Path, default=runtime("st1-040-content-probes.json"))
    parser.add_argument("--families", type=int, default=3)
    parser.add_argument("--members-per-family", type=int, default=3)
    parser.add_argument("--min-filename-year", type=int, help="metadata-only discovery filter; never a source-date assertion")
    parser.add_argument("--eligible-index", type=int, default=0, help="zero-based family after deterministic eligibility filters")
    args = parser.parse_args()
    if not args.root.is_dir():
        raise SystemExit("approved pilot root is unavailable")
    if args.family_manifest:
        manifest = json.loads(args.family_manifest.read_text(encoding="utf-8"))
        signature = manifest["selection_signature"]
        selected = [{
            "relative_locator": manifest["relative_locator"],
            "document_count": signature["document_count"],
            "locally_probeable_document_count": signature["probeable_document_count"],
            "negative_context": False,
            "files": manifest["files"],
        }]
    else:
        ranking = json.loads(args.ranking.read_text(encoding="utf-8"))
        selected = [
            family for family in ranking["top_families"]
            if 1 <= family["document_count"] <= 60 and family["locally_probeable_document_count"] > 0 and not family["negative_context"]
        ]
        if args.min_filename_year:
            pattern = re.compile(rf"(?<!\d){args.min_filename_year}(?!\d)")
            selected = [family for family in selected if any(pattern.search(file["filename"]) for file in family["files"])]
        selected = selected[args.eligible_index:args.eligible_index + args.families]
    if not selected:
        raise SystemExit("no bounded eligible source family found")
    probes = []
    for family in selected:
        candidates = sorted(
            (file for file in family["files"] if file["extension"] in ALLOWED),
            key=lambda file: (file.get("modified_utc") or "", file["relative_locator"].casefold()),
            reverse=True,
        )[:args.members_per_family]
        for file in candidates:
            source = args.root / file["relative_locator"]
            item = {"family_locator": family["relative_locator"], "file": file, "probe_status": "unknown"}
            try:
                result = extract(source)
                text = result.pop("text")
                item.update(result)
                item["keyword_signals"] = evidence_signals(text)
                item["sample_text"] = text
                item["probe_status"] = "extracted"
            except Exception as exc:  # diagnostic only; do not expose source details in Git
                item["probe_status"] = "error"
                item["error_type"] = type(exc).__name__
            probes.append(item)
    output = {
        "schema_version": "st1-040-content-probes-v1",
        "generated_utc": datetime.now(UTC).isoformat(),
        "metadata_ranking": {"families_considered": len(selected), "members_per_family": args.members_per_family},
        "probes": probes,
        "boundaries": {"read_only": True, "external_model_use": False, "platform_persistence": False, "automatic_certification": False, "raw_output_outside_git": True},
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "families_probed": len(selected),
        "members_attempted": len(probes),
        "members_extracted": sum(item["probe_status"] == "extracted" for item in probes),
        "members_with_errors": sum(item["probe_status"] == "error" for item in probes),
        "output_outside_git": True,
    }, separators=(",", ":")))


if __name__ == "__main__":
    main()
